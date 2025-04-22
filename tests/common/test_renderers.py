from io import BytesIO

import pandas as pd
from docx import Document
from rest_framework.response import Response

from bmds_ui.common import renderers
from pybmds.reporting.styling import Report


class TestDocxRenderer:
    def test_success(self):
        report = Report.build_default()
        report.document.add_heading("Test123")
        file = BytesIO()
        report.document.save(file)
        file = renderers.BinaryFile(data=file, filename="test")

        resp = Response()
        data = renderers.DocxRenderer().render(data=file, renderer_context={"response": resp})
        d2 = Document(BytesIO(data))
        assert resp["Content-Disposition"] == 'attachment; filename="test.docx"'
        assert d2.paragraphs[0].text == "Test123"

    def test_error(self):
        resp = Response()
        error = {"test": "here"}
        b = renderers.DocxRenderer().render(data=error, renderer_context={"response": resp})
        d2 = Document(BytesIO(b))
        assert resp["Content-Disposition"] == 'attachment; filename="error.docx"'
        assert d2.paragraphs[0].text == "An error occurred"
        assert d2.paragraphs[1].text == '{\n  "test": "here"\n}'


class TestExcelRenderer:
    def test_success(self):
        df = pd.DataFrame(columns=["a", "b"], data=[[1, 2], [4, 5]])
        f = BytesIO()
        with pd.ExcelWriter(f) as writer:
            df.to_excel(writer, index=False)
        file = renderers.BinaryFile(data=f, filename="test")

        resp = Response()
        data = renderers.XlsxRenderer().render(data=file, renderer_context={"response": resp})

        df2 = pd.read_excel(BytesIO(data))
        assert df2.to_dict(orient="records") == [{"a": 1, "b": 2}, {"a": 4, "b": 5}]
        assert resp["Content-Disposition"] == 'attachment; filename="test.xlsx"'

    def test_error(self):
        resp = Response()
        error = {"test": "here"}
        resp = Response()
        data = renderers.XlsxRenderer().render(data=error, renderer_context={"response": resp})
        df2 = pd.read_excel(BytesIO(data))
        assert df2.Status[0] == '"{\\n  \\"test\\": \\"here\\"\\n}"'
