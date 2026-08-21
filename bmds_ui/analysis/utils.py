import base64
import io
import re

from django.conf import settings
from django.utils.timezone import now
from docx.shared import Inches
from PIL import Image

from pybmds.utils import get_version

from .. import __version__


def get_citation() -> str:
    """
    Return a citation for the software.
    """
    year = "20" + __version__[:2]
    accessed = now().strftime("%B %d, %Y")
    version = get_version()
    application = "BMDS Desktop" if settings.IS_DESKTOP else "BMDS Online"
    uri = "https://pypi.org/project/bmds-ui/" if settings.IS_DESKTOP else settings.WEBSITE_URI
    return f"U.S. Environmental Protection Agency. ({year}). {application} ({__version__}; pybmds {version.python}; bmdscore {version.dll}) [Software]. Available from {uri}. Accessed {accessed}."


re_hex_color = re.compile("^#(?:[0-9a-fA-F]{3}){1,2}$")


def fig_to_png_b64(fig, dpi: int = 100, tight: bool = False) -> str:
    buf = io.BytesIO()
    kwargs = {"format": "png", "dpi": dpi}
    if tight:
        kwargs["bbox_inches"] = "tight"
    fig.savefig(buf, **kwargs)
    buf.seek(0)

    img = Image.open(buf)
    # quantize to a palette if the image doesn't need full RGBA
    if img.mode in ("RGBA", "RGB"):
        img = img.convert("P", palette=Image.ADAPTIVE, colors=256)

    out = io.BytesIO()
    img.save(out, format="PNG", optimize=True, compress_level=9)
    out.seek(0)
    return base64.b64encode(out.getvalue()).decode("utf-8")


def add_png_b64_to_docx(
    document, b64_png: str, width_in: float | None = 6.0, height_in: float | None = None
):
    """
    Add a base64-encoded PNG to a python-docx Document.
    If width_in or height_in is provided, python-docx will scale proportionally.
    """
    # Handle possible data URL prefix
    if b64_png.startswith("data:image"):
        b64_png = b64_png.split(",", 1)[-1]

    img_bytes = base64.b64decode(b64_png)
    stream = io.BytesIO(img_bytes)

    if width_in is not None and height_in is not None:
        document.add_picture(stream, width=Inches(width_in), height=Inches(height_in))
    elif width_in is not None:
        document.add_picture(stream, width=Inches(width_in))
    elif height_in is not None:
        document.add_picture(stream, height=Inches(height_in))
    else:
        document.add_picture(stream)  # native size
